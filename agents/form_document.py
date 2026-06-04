"""
Đọc / điền file Word (.docx, .doc) — luôn ghi ra bản sao, không sửa file gốc trong docs/.
Điền ưu tiên giữ format run (python-docx); .doc fallback win32com.
"""

from __future__ import annotations

import logging
import re
import shutil
import subprocess
from copy import deepcopy
from pathlib import Path

from agents.form_field_normalize import date_parts, normalize_field_value
from agents.form_field_schema import all_labels_for, resolve_fill_template

log = logging.getLogger(__name__)


def extract_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def extract_doc_text(path: Path) -> str:
    try:
        result = subprocess.run(
            ["antiword", str(path)],
            capture_output=True,
            text=True,
            timeout=30,
            encoding="utf-8",
            errors="ignore",
        )
        text = (result.stdout or "").strip()
        if len(text) > 40:
            return text
    except Exception as e:
        log.debug(f"antiword unavailable for {path.name}: {e}")

    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(path.resolve()))
        text = doc.Content.Text.replace("\r", "\n")
        doc.Close(False)
        word.Quit()
        return text.strip()
    except Exception as e:
        log.warning(f"win32com extract failed {path.name}: {e}")

    raw = path.read_bytes()
    text = raw.decode("utf-16-le", errors="ignore")
    lines = [
        ln.strip()
        for ln in text.split("\n")
        if 4 < len(ln.strip()) < 220 and re.search(r"[A-Za-zÀ-ỹ]", ln)
    ]
    return "\n".join(lines[:120])


def extract_form_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_docx_text(path)
    if ext == ".doc":
        return extract_doc_text(path)
    raise ValueError(f"Unsupported form format: {ext}")


_NEXT_FIELD_MARKERS = (
    "đến ngày", "đến ngày ", "số điện thoại", "giới tính", "mã sinh viên", "ngành:", "lớp:",
)


def _slot_is_empty(text: str, label: str) -> bool:
    for sep in (":", "：", ""):
        needle = label + sep if sep else label
        if needle not in text:
            continue
        idx = text.index(needle) + len(needle)
        rest = text[idx:]
        for part in rest.split("\t"):
            chunk = part.strip()
            if not chunk:
                continue
            low = chunk.lower()
            if any(low.startswith(m) for m in _NEXT_FIELD_MARKERS):
                return True
            if chunk.endswith(":") or chunk.endswith("："):
                return True
            if re.fullmatch(r"[_…\.\s]+", chunk):
                continue
            return len(chunk) < 2
        return True
    return False


def _fill_label_in_text(text: str, label: str, value: str) -> str:
    if not value or label not in text:
        return text
    if not _slot_is_empty(text, label):
        return text

    for sep in (":", "："):
        needle = label + sep
        if needle not in text:
            continue
        idx = text.index(needle) + len(needle)
        rest = text[idx:]
        return text[:idx] + "\t" + value + rest

    return text.replace(label, f"{label}: {value}", 1)


def _labels_for_field(
    key: str,
    primary_label: str,
    extra_labels: dict[str, list[str]] | None = None,
) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []

    def add(lbl: str) -> None:
        lbl = lbl.strip()
        if lbl and lbl not in seen:
            seen.add(lbl)
            out.append(lbl)

    add(primary_label)
    if extra_labels:
        for alias in extra_labels.get(key, []):
            add(alias)
    return out


def _fill_paragraph_text(
    text: str,
    field_labels: dict[str, str],
    answers: dict[str, str],
    extra_labels: dict[str, list[str]] | None = None,
) -> str:
    jobs: list[tuple[int, str, str]] = []
    for key, primary in field_labels.items():
        val = (answers.get(key) or "").strip()
        if not val:
            continue
        for label in _labels_for_field(key, primary, extra_labels):
            pos = text.find(label)
            if pos >= 0:
                jobs.append((pos, label, val))
                break

    jobs.sort(key=lambda x: (-len(x[1]), x[0]))
    new_t = text
    for _, label, val in jobs:
        new_t = _fill_label_in_text(new_t, label, val)
    return new_t


def _char_offset_to_run_index(para, char_pos: int) -> tuple[int, int]:
    """Trả về (run_index, offset_trong_run)."""
    pos = 0
    for i, run in enumerate(para.runs):
        n = len(run.text)
        if pos + n >= char_pos:
            return i, char_pos - pos
        pos += n
    return max(0, len(para.runs) - 1), 0


def _insert_run_after(para, after_run, text: str) -> None:
    from docx.text.run import Run

    new_r = deepcopy(after_run._r)
    after_run._r.addnext(new_r)
    new_run = Run(new_r, para)
    new_run.text = text
    for attr in ("bold", "italic", "underline", "font"):
        try:
            setattr(new_run, attr, getattr(after_run, attr))
        except Exception:
            pass


def _fill_paragraph_preserving_runs(
    para,
    field_labels: dict[str, str],
    answers: dict[str, str],
    extra_labels: dict[str, list[str]] | None,
) -> bool:
    """Điền tại chỗ nhãn, cố giữ run gốc."""
    full = para.text
    if not full.strip():
        return False

    jobs: list[tuple[int, str, str]] = []
    for key, primary in field_labels.items():
        val = (answers.get(key) or "").strip()
        if not val:
            continue
        for label in _labels_for_field(key, primary, extra_labels):
            for sep in (":", "：", ""):
                needle = label + sep if sep else label
                pos = full.find(needle)
                if pos < 0:
                    continue
                insert_at = pos + len(needle)
                if sep and not _slot_is_empty(full, label):
                    break
                if sep == "" and not _slot_is_empty(full, label):
                    continue
                jobs.append((insert_at, label, val))
                break
            if jobs and jobs[-1][1] == label:
                break

    if not jobs:
        return False

    jobs.sort(key=lambda x: (-x[0], -len(x[1])))
    changed = False
    for insert_at, label, val in jobs:
        current = para.text
        if label not in current:
            continue
        for sep in (":", "："):
            needle = label + sep
            if needle in current:
                idx = current.index(needle) + len(needle)
                if not _slot_is_empty(current, label):
                    break
                ri, off = _char_offset_to_run_index(para, idx)
                if not para.runs:
                    break
                run = para.runs[ri]
                prefix = run.text[:off]
                suffix = run.text[off:]
                run.text = prefix
                insert_text = "\t" + val
                if suffix.strip():
                    _insert_run_after(para, run, insert_text + suffix)
                else:
                    _insert_run_after(para, run, insert_text)
                changed = True
                break
    return changed


def _apply_date_triplets(
    para,
    key: str,
    raw_value: str,
    labels: list[str],
) -> bool:
    """Một số mẫu có «ngày … tháng … năm» — chèn dd/mm/yyyy sau nhãn ngày sinh."""
    parts = date_parts(raw_value)
    if not parts:
        return False
    d, mo, y = parts
    combined = f"{d}/{mo}/{y}"
    return _fill_paragraph_preserving_runs(
        para,
        {key: labels[0] if labels else key},
        {key: combined},
        {key: labels},
    )


def normalize_answers(
    answers: dict[str, str],
    field_types: dict[str, str],
) -> dict[str, str]:
    out: dict[str, str] = {}
    for key, raw in answers.items():
        ft = field_types.get(key, "text")
        out[key] = normalize_field_value(ft, raw)
    return out


def fill_docx(
    template: Path,
    output: Path,
    field_labels: dict[str, str],
    answers: dict[str, str],
    *,
    catalog_filename: str = "",
    field_types: dict[str, str] | None = None,
) -> list[str]:
    """
    Điền docx; trả về danh sách key không khớp nhãn trên mẫu.
    """
    from docx import Document

    ft = field_types or {}
    norm_answers = normalize_answers(answers, ft)
    extra = all_labels_for(catalog_filename) if catalog_filename else {}

    shutil.copy2(template, output)
    doc = Document(str(output))
    unfilled: set[str] = set(norm_answers.keys())

    def process_para(para) -> None:
        nonlocal unfilled
        before = para.text
        used_keys: set[str] = set()

        if _fill_paragraph_preserving_runs(para, field_labels, norm_answers, extra):
            after = para.text
            for key in norm_answers:
                if norm_answers[key] and norm_answers[key] in after and norm_answers[key] not in before:
                    used_keys.add(key)
        else:
            new_t = _fill_paragraph_text(before, field_labels, norm_answers, extra)
            if new_t != before:
                para.text = new_t
                for key, val in norm_answers.items():
                    if val and val in new_t and val not in before:
                        used_keys.add(key)

        for key in used_keys:
            unfilled.discard(key)
            if ft.get(key) == "date":
                labels = extra.get(key, [field_labels.get(key, "")])
                _apply_date_triplets(para, key, norm_answers[key], labels)

    for para in doc.paragraphs:
        process_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_para(para)

    doc.save(str(output))
    return sorted(unfilled)


def fill_doc_win32(
    template: Path,
    output: Path,
    field_labels: dict[str, str],
    answers: dict[str, str],
    *,
    catalog_filename: str = "",
    field_types: dict[str, str] | None = None,
) -> list[str]:
    import win32com.client  # type: ignore

    ft = field_types or {}
    norm_answers = normalize_answers(answers, ft)
    extra = all_labels_for(catalog_filename) if catalog_filename else {}

    shutil.copy2(template, output)
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(str(output.resolve()))
    filled_keys: set[str] = set()

    for key, primary in field_labels.items():
        value = (norm_answers.get(key) or "").strip()
        if not value:
            continue
        for label in _labels_for_field(key, primary, extra):
            for find_text in (f"{label}:", f"{label}：", label):
                rng = doc.Content
                f = rng.Find
                f.ClearFormatting()
                f.Text = find_text
                f.Forward = True
                f.Wrap = 0
                if f.Execute():
                    rng.SetRange(rng.End, rng.End)
                    rng.InsertAfter(f"\t{value}")
                    filled_keys.add(key)
                    break
            else:
                continue
            break

    doc.Save()
    doc.Close(False)
    word.Quit()
    return sorted(k for k in norm_answers if k not in filled_keys and norm_answers[k])


def fill_form_copy(
    template: Path,
    output: Path,
    field_labels: dict[str, str],
    answers: dict[str, str],
    *,
    catalog_filename: str = "",
    field_types: dict[str, str] | None = None,
) -> list[str]:
    """Điền bản sao; trả về key chưa ghi được."""
    output.parent.mkdir(parents=True, exist_ok=True)
    ext = template.suffix.lower()
    kwargs = {
        "catalog_filename": catalog_filename,
        "field_types": field_types,
    }
    if ext == ".docx":
        return fill_docx(template, output, field_labels, answers, **kwargs)
    if ext == ".doc":
        return fill_doc_win32(template, output, field_labels, answers, **kwargs)
    raise ValueError(f"Unsupported: {ext}")


def resolve_template_path(bieu_mau_dir: Path, catalog_filename: str) -> Path:
    """Chọn file mẫu để điền (.docx ưu tiên)."""
    fill_name = resolve_fill_template(catalog_filename)
    path = bieu_mau_dir / fill_name
    if path.is_file():
        return path
    legacy = bieu_mau_dir / catalog_filename
    if legacy.is_file():
        return legacy
    raise FileNotFoundError(catalog_filename)
